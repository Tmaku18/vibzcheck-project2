"""Generate Makuvaza_Tanaka_Vibzcheck_Project2_Proposal.docx.

Mirrors the python-docx generator pattern used for prior CW deliverables.
Run from the Project2 directory:

    python make_vibzcheck_proposal_doc.py
"""

from __future__ import annotations

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


STUDENT_NAME = "Tanaka Makuvaza"
STUDENT_ID = "002252191"
COURSE = "CSC 4360/6370 - Mobile App Development"
TERM = "Spring 2026"
CRN = "13598"
PROPOSAL_DATE = "April 20, 2026"
TEAM_NAME = "Vibzcheck - Solo Team (Tanaka Makuvaza)"
GITHUB_URL = "https://github.com/Tmaku18/vibzcheck-project2"
WORD_FILENAME = "Makuvaza_Tanaka_Vibzcheck_Project2_Proposal.docx"


def add_centered_run(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    doc.add_paragraph(text, style=style)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    table.autofit = True


def add_table(doc, header, rows, style="Light Grid Accent 1"):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = style
    for j, label in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.rows[i].cells[j].text = value
    return table


def build_doc():
    doc = Document()

    # ----- Title block -----
    add_centered_run(
        doc, "Vibzcheck - Collaborative Music App", bold=True, size=18, space_after=2
    )
    add_centered_run(
        doc,
        "Project 2 Proposal | Flutter & Firebase Final Group Project",
        bold=True,
        size=12,
        space_after=2,
    )
    add_centered_run(
        doc, f"{COURSE} | {TERM} | CRN {CRN}", size=11, space_after=2
    )
    add_centered_run(
        doc,
        f"{STUDENT_NAME} | Student ID {STUDENT_ID} | Graduate (M.S. Computer Science)",
        size=11,
        space_after=2,
    )
    add_centered_run(doc, f"Proposal Date: {PROPOSAL_DATE}", size=11, space_after=18)

    # ----- 1. Submission Metadata -----
    add_heading(doc, "1. Submission Metadata")
    add_kv_table(
        doc,
        [
            ("Course", COURSE),
            ("Section / CRN", CRN),
            ("Term", TERM),
            ("Team Name", TEAM_NAME),
            ("Proposal Date", PROPOSAL_DATE),
            ("GitHub Repository URL", GITHUB_URL),
            ("Word Document Filename", WORD_FILENAME),
            ("Degree Level", "Graduate (Master's) - includes advanced extension"),
        ],
    )
    doc.add_paragraph()

    # ----- 2. Project Overview -----
    add_heading(doc, "2. Project Overview")
    add_paragraph(
        doc,
        "Vibzcheck is a collaborative music streaming application that transforms the traditional "
        "playlist experience into a dynamic, group-oriented platform. By combining real-time voting, "
        "Spotify Web API integration, and Firebase backend services, Vibzcheck enables groups of friends "
        "to collectively curate playlists, share vibes, and discover music together.",
    )
    add_paragraph(
        doc,
        "Core problem: current music apps treat listening as a solitary experience even when people are "
        "together. Vibzcheck closes that gap by making music curation a collaborative, real-time activity "
        "where every member shapes the queue. Users add tracks to a shared queue, vote on what plays next, "
        "tag songs with mood and genre, and chat about the session in real time.",
    )
    add_paragraph(
        doc,
        "Target users: friend groups at parties, study sessions, road trips, dorms, and small social "
        "gatherings who want a shared, transparent way to decide what plays next without one person "
        "monopolizing the aux cable.",
    )

    # ----- 3. Team Members -----
    add_heading(doc, "3. Team Members")
    add_table(
        doc,
        header=["#", "Member Name", "Student ID", "Role / Responsibility", "Signed Date"],
        rows=[
            (
                "1",
                STUDENT_NAME,
                STUDENT_ID,
                "UI / Backend / Firebase / Testing / Documentation (sole owner)",
                "2026-04-20",
            )
        ],
    )
    add_paragraph(
        doc,
        "Note: Vibzcheck is being delivered as an explicit one-person team. All implementation, "
        "testing, documentation, and presentation responsibilities are owned by the sole member listed "
        "above.",
    )

    # ----- 4. Statement of Commitment -----
    add_heading(doc, "4. Statement of Commitment")
    add_bullets(
        doc,
        [
            "Individual contributions are documented through frequent, granular Git commits with meaningful messages.",
            "The sole team member can explain every implementation section, architectural decision, and Firebase rule.",
            "Code quality, automated and manual testing, and documentation will meet the project standards described in the Project 2 Submission Requirements.",
            "The proposal deadline (April 20, 2026, 11:59 PM) and the final delivery deadline (May 3, 2026, 11:59 PM) are firm and will be met.",
            "A signed statement PDF is attached to the submitted Word document.",
        ],
    )

    # ----- 5. Project Objectives -----
    add_heading(doc, "5. Project Objectives")

    add_heading(doc, "5.1 Objective 1 - User Identity & Authentication", level=2)
    add_paragraph(
        doc,
        "Implement secure user authentication via Firebase Authentication using email/password "
        "(with optional Google sign-in). Each user has a unique profile, persistent session, and "
        "auth-scoped access to sessions they own or have joined.",
    )

    add_heading(doc, "5.2 Objective 2 - Real-Time Playlist Management", level=2)
    add_paragraph(
        doc,
        "Build Firestore real-time collections to sync playlists, queue state, and voting data across "
        "all members of a session. Use Firestore transactions to keep vote counts and queue ordering "
        "consistent under concurrent updates.",
    )

    add_heading(doc, "5.3 Objective 3 - Spotify API Integration", level=2)
    add_paragraph(
        doc,
        "Integrate the Spotify Web API through Firebase Cloud Functions to search for tracks, fetch "
        "metadata, audio features, and 30-second previews. Spotify client credentials never leave the "
        "server; the mobile client only sees enriched, sanitized track data.",
    )

    add_heading(doc, "5.4 Objective 4 - Real-Time Collaboration", level=2)
    add_paragraph(
        doc,
        "Enable live group interactions through Firestore listeners: voting on the next track, tagging "
        "songs with mood and genre, exchanging chat messages, and viewing collective preferences that "
        "shape recommendations.",
    )

    add_heading(doc, "5.5 Objective 5 - AI Must-Solve Helper (Required)", level=2)
    add_paragraph(
        doc,
        "Implement the required AI must-solve challenge: a Firestore-backed mood + vote helper that "
        "suggests the next 3 songs using transparent if/then scoring rules and exposes a manual "
        "override control. Suggestions and the rule snapshot used to produce them are persisted so the "
        "logic can be inspected during the demo and Q&A.",
    )

    add_heading(doc, "5.6 Objective 6 - Graduate-Level Extension", level=2)
    add_paragraph(
        doc,
        "Implement a fairness-based playlist ranking module that balances individual preferences with "
        "group dynamics. Tracks are scored using vote weight, recent-play decay, and per-member "
        "participation so that no single member dominates the queue. Each ranked suggestion exposes the "
        "factors that produced its score so the algorithm is defensible at code level.",
    )

    # ----- 6. Functional Requirements & Phases -----
    add_heading(doc, "6. Functional Requirements and Phase Breakdown")
    add_table(
        doc,
        header=["Phase", "Window", "Functional Requirements"],
        rows=[
            (
                "Phase 1",
                "Apr 14 - Apr 20",
                "Firebase Authentication, Firestore user profiles, security rules baseline, app shell "
                "and navigation, proposal package.",
            ),
            (
                "Phase 2",
                "Apr 21 - Apr 24",
                "Collaborative session/playlist system, transaction-backed voting, Spotify bridging via "
                "Cloud Functions, basic chat scaffold.",
            ),
            (
                "Phase 3",
                "Apr 25 - Apr 29",
                "Mood/genre tagging, AI must-solve helper, fairness ranking module, FCM notifications, "
                "Storage-backed avatars/album art, animations and polish.",
            ),
            (
                "Phase 4",
                "Apr 30 - May 3",
                "Performance optimization, edge case handling, automated/manual testing, evidence "
                "capture, slide deck, demo video, APK, final delivery.",
            ),
        ],
    )

    # ----- 7. Technology Stack & Firebase Implementation -----
    add_heading(doc, "7. Technology Stack and Firebase Implementation")
    add_paragraph(
        doc,
        "Vibzcheck is built on Flutter (Dart) for the cross-platform mobile client and Google Firebase "
        "for the backend. Firebase eliminates infrastructure overhead, provides built-in authentication "
        "and real-time synchronization, scales automatically, and integrates natively with Flutter via "
        "FlutterFire. Cloud Functions are used as a secure server-side bridge for the Spotify Web API "
        "so that client credentials are never shipped with the mobile app.",
    )

    add_heading(doc, "7.1 Firebase Authentication", level=2)
    add_paragraph(
        doc,
        "Email/password sign-up and sign-in with optional Google provider. Each authenticated session "
        "owner becomes the auth-scoped subject in Firestore security rules.",
    )

    add_heading(doc, "7.2 Cloud Firestore", level=2)
    add_paragraph(
        doc,
        "Real-time NoSQL database for users, sessions, members, queue, votes, chat messages, mood "
        "tags, and recommendation snapshots. Atomic transactions keep vote counts and queue ordering "
        "consistent under concurrent edits.",
    )

    add_heading(doc, "7.3 Firebase Cloud Functions", level=2)
    add_paragraph(
        doc,
        "Serverless backend logic for the Spotify Web API bridge, recommendation/fairness scoring "
        "snapshots, and notification triggers. Spotify client ID and secret live only in the function "
        "environment.",
    )

    add_heading(doc, "7.4 Firebase Storage", level=2)
    add_paragraph(
        doc,
        "Stores user avatars and cached album/playlist artwork with auth-scoped access rules. Storage "
        "paths are persisted as references on the related Firestore documents.",
    )

    add_heading(doc, "7.5 Firebase Cloud Messaging (FCM)", level=2)
    add_paragraph(
        doc,
        "Push notifications when a user is invited to a session, when a new track is queued, or when a "
        "voting round opens. Foreground, background, and terminated app states are handled distinctly.",
    )

    add_heading(doc, "7.6 Flutter (Dart) Frontend", level=2)
    add_paragraph(
        doc,
        "Cross-platform UI built with Flutter widgets and the FlutterFire packages "
        "(firebase_core, firebase_auth, cloud_firestore, firebase_storage, firebase_messaging, "
        "cloud_functions).",
    )

    # ----- 8. Architecture & Data Model -----
    add_heading(doc, "8. Architecture and Data Model Outline")
    add_paragraph(
        doc,
        "High-level data flow: the Flutter client talks directly to Authentication, Firestore, Storage, "
        "and Cloud Messaging. All Spotify traffic is brokered by Cloud Functions, which enrich Spotify "
        "responses and persist sanitized track data into Firestore so the client only ever reads "
        "Firebase-owned documents.",
    )
    add_heading(doc, "8.1 Firestore Collection Layout", level=2)
    add_bullets(
        doc,
        [
            "users/{uid} - profile, avatar storage path, favorite genres, FCM token, notification preferences.",
            "sessions/{sessionId} - owner uid, title, status, currentTrackId, mood summary, timestamps.",
            "sessions/{sessionId}/members/{uid} - joinedAt, role, presence, lastVoteAt.",
            "sessions/{sessionId}/queue/{trackId} - track metadata, addedBy, voteScore, mood tags, order, artwork storage path.",
            "sessions/{sessionId}/messages/{messageId} - sender uid, text, emoji reactions, createdAt.",
            "sessions/{sessionId}/suggestions/{snapshotId} - ranked next-song suggestions plus the explainable factors that produced each score.",
        ],
    )
    add_heading(doc, "8.2 Security Considerations", level=2)
    add_bullets(
        doc,
        [
            "All Firestore reads and writes are gated by request.auth.uid; unauthenticated requests are rejected.",
            "Session subcollections (queue, members, messages, suggestions) require the requester to be a member of the parent session.",
            "Storage paths are scoped per uid (avatars) or per session (artwork) and validated through Storage security rules.",
            "Spotify credentials are stored only in Cloud Functions environment configuration, never shipped with the client.",
            "Rule changes are validated using the Firestore Rules emulator and a small library of allow/deny test cases before deploy.",
        ],
    )

    # ----- 9. Wireframe & UI Design Plan -----
    add_heading(doc, "9. Wireframe and UI Design Plan")
    add_paragraph(
        doc,
        "The UI is structured around a small, focused set of screens that follow the user's natural "
        "flow from sign-in to collaborative listening to reflection. Wireframe artifacts (Figma / "
        "Draw.io exports) will be linked alongside this proposal at submission time.",
    )
    add_table(
        doc,
        header=["Screen", "Primary Responsibility"],
        rows=[
            ("Sign In / Sign Up", "Firebase Authentication entry point and initial profile creation."),
            ("Home / Session Lobby", "Create or join a session; show active sessions and member presence."),
            ("Session / Queue", "Shared queue, real-time voting, mood/genre tags, current track display."),
            ("Track Search", "Spotify-backed search via Cloud Functions; add tracks to the active session queue."),
            ("Chat", "Real-time message subcollection with emoji reactions tied to the active session."),
            ("Suggestions / Insights", "Rule-based next-3-song helper plus the graduate fairness-ranking view with explainable factors."),
            ("Profile", "Avatar (Storage), favorite genres, listening history summary."),
            ("Settings", "Notification preferences (FCM topic toggles), privacy controls, sign out."),
        ],
    )
    add_heading(doc, "9.1 Navigation Flow", level=2)
    add_paragraph(
        doc,
        "User journey: Sign In -> Home Lobby -> Create or Join Session -> Session/Queue (vote, tag, "
        "chat) -> Suggestions/Insights -> Profile and Settings. Bottom navigation gives quick access to "
        "the active session, search, suggestions, and profile. Firestore listeners refresh the queue, "
        "votes, and chat without manual reloads.",
    )
    add_heading(doc, "9.2 Design Principles", level=2)
    add_bullets(
        doc,
        [
            "Clear visual hierarchy: primary actions (vote, add track) are dominant; secondary options live in menus.",
            "Real-time feedback: instant vote updates, live presence indicators, queue position changes.",
            "Mood-based cues: color and iconography reinforce the collective mood/genre tags.",
            "Accessibility: high contrast palette, readable font sizes, large tap targets.",
            "Responsive design: portrait/landscape support and scaling for typical 5\"-7\" device sizes.",
        ],
    )
    add_paragraph(doc, "Wireframe link: [ATTACH FIGMA/DRAW.IO LINK]")

    # ----- 10. Testing & Evidence Plan -----
    add_heading(doc, "10. Testing and Evidence Plan")
    add_table(
        doc,
        header=["Layer", "What is tested", "Evidence captured"],
        rows=[
            (
                "Unit",
                "Vote scoring math, fairness ranking factors, AI helper rule evaluation.",
                "Test files under test/ with passing CI output and screenshots.",
            ),
            (
                "Widget",
                "Critical UI flows: session creation, vote tap, queue render, suggestion override.",
                "Widget test outputs and golden screenshots.",
            ),
            (
                "Integration / Manual",
                "End-to-end session: sign in, create session, add track, multi-device vote, chat, FCM delivery.",
                "Recorded walkthrough plus per-flow screenshot evidence.",
            ),
            (
                "Security Rules",
                "Allow/deny matrix for users, sessions, queue, messages, storage paths.",
                "Firestore Rules emulator test output and one captured blocked-write proof.",
            ),
            (
                "Failure Modes",
                "Network drop, auth expiry, missing FCM token, Cloud Function error.",
                "Captured error-state UI and recovery path notes.",
            ),
        ],
    )

    # ----- 11. Development Timeline -----
    add_heading(doc, "11. Development Timeline (Apr 14 - May 3)")
    add_table(
        doc,
        header=["Window", "Focus", "Key Deliverables"],
        rows=[
            (
                "Apr 14 - 20",
                "Foundation and proposal",
                "Firebase project, Auth, profile flow, schema and rules baseline, app shell, "
                "submitted proposal package.",
            ),
            (
                "Apr 21 - 24",
                "Core collaborative workflow",
                "Sessions, shared queue, transaction-safe voting, Spotify bridge via Cloud Functions.",
            ),
            (
                "Apr 25 - 29",
                "Advanced features and polish",
                "Chat, mood/genre tagging, FCM, AI must-solve helper, fairness ranking, animations.",
            ),
            (
                "Apr 30 - May 3",
                "Hardening and submission",
                "Tests, optimization, edge case handling, evidence capture, slides, demo video, APK, "
                "final upload by May 3 11:59 PM.",
            ),
        ],
    )
    add_paragraph(
        doc,
        "Milestone discipline: granular Git commits per file change, daily self-review against the "
        "rubric criteria, and progressive evidence capture so the final demo and slides are anchored "
        "to commits, screenshots, and rule tests rather than reconstructed at the end.",
    )

    # ----- 12. Graduate-Level Rationale -----
    add_heading(doc, "12. Graduate-Level Rationale")
    add_paragraph(
        doc,
        "The graduate fairness-ranking module is justified because Vibzcheck's core value proposition "
        "is collective, not individual, music selection. Naive vote ordering allows the most active "
        "user to dominate the queue, which violates the product premise. The fairness ranking weights "
        "vote score, recent-play decay, and per-member participation so that under-represented voices "
        "surface naturally. Each ranked track exposes its component factors in the UI so the algorithm "
        "can be defended at code level during Q&A.",
    )
    add_paragraph(
        doc,
        "The required AI must-solve helper (next-3-song suggestions from mood + vote state) is "
        "implemented as a transparent rule pipeline rather than an opaque model so it satisfies the "
        "course's responsible-AI guidance: rules are explicit, suggestions can be manually overridden, "
        "and every suggestion snapshot is persisted in Firestore for inspection.",
    )

    # ----- 13. Submission Information & Deliverables -----
    add_heading(doc, "13. Submission Information and Deliverables")
    add_kv_table(
        doc,
        [
            ("GitHub Repository", GITHUB_URL),
            ("Repository Visibility", "Public (no instructor invite required)"),
            ("Word Document", WORD_FILENAME),
            ("Signed Statement PDF", "Attached to the Word document at submission time"),
            ("APK", "Built and uploaded for the final May 3 delivery"),
            ("Slide Deck", "PowerPoint deck used in the final 20-25 minute presentation"),
            (
                "Presentation Video",
                "Recorded demo uploaded to YouTube or Google Drive; link submitted to iCollege Dropbox",
            ),
        ],
    )

    # ----- 14. Quality Checklist -----
    add_heading(doc, "14. Quality Checklist")
    add_bullets(
        doc,
        [
            "Team member name, ID, and role clearly documented (solo team noted explicitly).",
            "Project concept is original and solves a real, defined problem.",
            "All five Firebase services (Auth, Firestore, Storage, FCM, Cloud Functions) are explained and planned.",
            "Wireframes show all core screens with clear navigation flows (artifacts to be attached).",
            "Phase breakdown is realistic and fits the Apr 14 - May 3 window.",
            "Git workflow strategy is defined: granular per-file commits with meaningful messages on a public repo.",
            "Testing approach (unit, widget, integration, security rules, failure modes) is documented.",
            "Sole team member understands and agrees to the proposal contents.",
            "Signed commitment statement PDF will be attached at upload time.",
            "Document is professionally formatted and free of spelling errors.",
        ],
    )

    return doc


def main():
    doc = build_doc()
    doc.save(WORD_FILENAME)
    print(f"Wrote {WORD_FILENAME}")


if __name__ == "__main__":
    main()
